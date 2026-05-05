import asyncio
import json
import re
import structlog
from concurrent.futures import ThreadPoolExecutor
from typing import Optional

import openai
from .prompts import (
    ENHANCE_SYSTEM, ENHANCE_USER,
    MAX_RESUME_CHARS, MAX_JD_CHARS,
)
from .schemas import (
    ContactInfo, EnhancedResume, ExperienceItem,
    EducationItem, TemplateStyle,
)

logger = structlog.get_logger(__name__)
_executor = ThreadPoolExecutor(max_workers=4, thread_name_prefix="doc-parser")
_cache: dict[str, EnhancedResume] = {}
_CACHE_MAX = 128

ACCEPTED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
}

IMAGE_MIME = {"image/jpeg", "image/jpg", "image/png", "image/gif", "image/webp"}


class ResumeEnhancerService:
    def __init__(self, api_key: str, model: str = "anthropic/claude-sonnet-4-5"):
        self._model = model
        self._client = openai.AsyncOpenAI(
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1",
        )

    async def enhance(
        self,
        file_bytes: bytes,
        filename: str,
        content_type: str,
        job_description: str,
        style: TemplateStyle = TemplateStyle.MODERN,
    ) -> EnhancedResume:
        
        if content_type not in ACCEPTED_MIME:
            raise ValueError(
                f"Unsupported file type. Please upload PDF, DOCX, DOC, or TXT."
            )

        resume_text = await self._extract_text(file_bytes, filename, content_type)
        jd_clean = _sanitize(job_description, MAX_JD_CHARS)

        # 1. LLM Call
        raw_output = await self._call_llm(
            ENHANCE_SYSTEM,
            ENHANCE_USER.format(resume_text=resume_text, job_description=jd_clean),
        )
        
        # 2. Parse & Validate
        try:
            return self._build_enhanced_resume(raw_output, style)
        except Exception as e:
            logger.error("Failed to build resume from LLM response", error=str(e))
            raise ValueError(f"Could not parse resume: {str(e)}")

    async def _extract_text(self, file_bytes: bytes, filename: str, content_type: str) -> str:
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(
            _executor, _extract_text_from_file, file_bytes, filename, content_type
        )
        if len(text.strip()) < 50:
            raise ValueError("File contains no readable text.")
        return _sanitize(text, MAX_RESUME_CHARS)

    async def _call_llm(self, system: str, user: str) -> str:
        response = await self._client.chat.completions.create(
            model=self._model,
            temperature=0.2,
            max_tokens=4096,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        )
        return response.choices[0].message.content

    def _build_enhanced_resume(self, raw: str, style: TemplateStyle) -> EnhancedResume:
        data = _parse_json(raw)
        
        return EnhancedResume(
            style=style,
            contact=ContactInfo(**data.get("contact", {})),
            summary=data.get("summary", ""),
            skills=data.get("skills", []),
            experience=data.get("experience", []),
            education=data.get("education", []),
            certifications=data.get("certifications", []),
            keywords_woven_in=data.get("keywords_woven_in", []),
            keywords_added=data.get("keywords_added", []),
            enhancement_notes=data.get("enhancement_notes", ""),
            sections_modified=data.get("sections_modified", []),
            sections_untouched=data.get("sections_untouched", []),
        )


# ── Helper Functions ──────────────────────────────────────────────────────────

def _sanitize(text: str, max_chars: int) -> str:
    return "".join(c for c in text if c.isprintable() or c in "\n\t ")[:max_chars]


def _parse_json(raw: str) -> dict:
    """Cleans markdown artifacts and parses JSON."""
    try:
        cleaned = re.sub(r"^```json\s*|\s*```$", "", raw.strip(), flags=re.MULTILINE)
        cleaned = re.sub(r"^```\s*|\s*```$", "", cleaned.strip(), flags=re.MULTILINE)
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        logger.error("JSON decode failed", error=str(e))
        raise ValueError("AI response was not valid JSON.")


def _extract_text_from_file(file_bytes: bytes, filename: str, content_type: str) -> str:
    """
    Extract text from PDF, DOCX, DOC, or TXT files.
    Tries multiple methods with fallbacks for maximum robustness.
    """
    text = ""
    
    # Try unstructured first (works for all types)
    try:
        text = _extract_with_unstructured(file_bytes, filename, content_type)
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"Unstructured extraction failed: {e}")
    
    # Try type-specific extraction
    if content_type == "text/plain":
        text = _extract_text_file(file_bytes)
    elif content_type == "application/pdf":
        text = _extract_pdf(file_bytes)
    elif content_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ]:
        text = _extract_docx_or_doc(file_bytes)
    
    if text.strip():
        return text
    
    # Final fallback: try to decode as text
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        raise ValueError("Could not extract text from file.")


def _extract_with_unstructured(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Extract using unstructured library (primary method)."""
    try:
        from unstructured.partition.auto import partition
        import io
        
        file_obj = io.BytesIO(file_bytes)
        elements = partition(
            file=file_obj,
            file_filename=filename,
            content_type=content_type,
        )
        return "\n".join([str(el) for el in elements])
    except Exception as e:
        raise ValueError(f"Unstructured partition failed: {str(e)}")


def _extract_text_file(file_bytes: bytes) -> str:
    """Extract from plain text files."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def _extract_pdf(file_bytes: bytes) -> str:
    """Extract from PDF files with multiple fallbacks."""
    text = ""
    
    # Try pypdf
    try:
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception:
                continue
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"pypdf extraction failed: {e}")
    
    # Try pdfplumber
    try:
        import pdfplumber
        import io
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception:
                    continue
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
    
    # Try PyPDF2
    try:
        from PyPDF2 import PdfReader
        import io
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception:
                continue
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
    
    return text


def _extract_docx_or_doc(file_bytes: bytes) -> str:
    """Extract from DOCX and DOC files with multiple fallbacks."""
    text = ""
    
    # Try python-docx (for DOCX)
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"python-docx extraction failed: {e}")
    
    # Try docx2txt (works for both DOCX and DOC)
    try:
        import docx2txt
        import io
        text = docx2txt.process(io.BytesIO(file_bytes))
        if text.strip():
            return text
    except Exception as e:
        logger.warning(f"docx2txt extraction failed: {e}")
    
    # Try doc2docx (for DOC files)
    try:
        from doc2docx import convert
        import tempfile
        import os
        
        with tempfile.TemporaryDirectory() as tmp_dir:
            input_path = os.path.join(tmp_dir, "input.doc")
            with open(input_path, "wb") as f:
                f.write(file_bytes)
            
            output_path = os.path.join(tmp_dir, "output.docx")
            convert(input_path, output_path)
            
            from docx import Document
            doc = Document(output_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            if text.strip():
                return text
    except Exception as e:
        logger.warning(f"doc2docx extraction failed: {e}")
    
    return text